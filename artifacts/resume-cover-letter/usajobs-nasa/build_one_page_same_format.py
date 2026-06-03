from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE = Path(__file__).resolve().parent
SOURCE = Path(
    r"C:\Users\thurs\Downloads\Derek Martin Linkedin General resume updated Feb 5 2026 (1).docx"
)
OUTPUT = BASE / "derek-martin-resume-usajobs-nasa-one-page.docx"


def set_run_font(run, *, size=None, name=None, bold=None, italic=None):
    if name:
        run.font.name = name
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_text(paragraph, text, *, bold=None, italic=None, size=None):
    was_bold = any(run.bold for run in paragraph.runs)
    was_italic = any(run.italic for run in paragraph.runs)
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(
        run,
        name="Times New Roman",
        size=size,
        bold=was_bold if bold is None else bold,
        italic=was_italic if italic is None else italic,
    )


def insert_paragraph_after(paragraph, text):
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    paragraph._p.addnext(new_para._p)
    if paragraph._p.pPr is not None:
        new_para._p.get_or_add_pPr().append(deepcopy(paragraph._p.pPr))
    run = new_para.add_run(text)
    set_run_font(run, name="Times New Roman")
    return new_para


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def compact(paragraph, *, before=0, after=0, line=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if line is not None:
        fmt.line_spacing = line


def main():
    doc = Document(SOURCE)

    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    # USAJOBS recommends 0.5-inch margins; this also gives Word pagination
    # enough room for the expanded federal resume details.
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    replacements = {
        "Mechanical Engineering student (aerospace focus, physics minor) with strength in space systems thinking and hands-on mechanical design. Built Python automation tools and currently developing Python control software (and supporting CAD/integration) for an autonomous piano-player capstone (target delivery June 2026). Designed and analyzed a high-safety-factor aluminum winch stand to lift a 150 lb motor ~50 ft. Seeking remote part-time engineering work through June 2027 and launch/space roles starting 2027.": (
            "Mechanical Engineering student (aerospace focus, physics minor) with space systems coursework, "
            "hands-on mechanical design, Python control software, embedded mechatronics, and high-reliability "
            "field experience."
        ),
        "Engineering Intern | Evensol Inc. - EDG-L (Landfill Gas Collection) | June 2024 - September 2024": (
            "Engineering Intern | Evensol Inc. - EDG-L | June 2024 - September 2024 | Hours/week: varied; average 26"
        ),
        "     Wildland Firefighter | USFS Sisters Ranger District | August 2022 - September 2023": (
            "Wildland Firefighter | USFS Sisters Ranger District | August 2022 - September 2023 | GS-03/GS-04 | Hours/week: 40-112"
        ),
        'Awarded "Hardest Worker" in Calculus (2021)': (
            'Awarded "Hardest Worker" in Calculus | 2021 and 2022'
        ),
        "Team Captain for Nordic Skiing": "Team Captain for Nordic Skiing | 2021 - 2022",
        "Medallion awarded by Bend Police Department for integrity and discipline (2021)": (
            "Bend Police Department medallion for integrity and discipline | 2020"
        ),
        "Autonomous Piano Player (Capstone, 2025–2026): Developing Python control software for an electromechanical system; supporting CAD/integration as needed.": (
            "Autonomous Piano Player (Capstone, January 2026 - June 2026): Project leader for a bench-scale embedded mechatronics system converting MIDI files into real-time actuator commands for 62 hardware channels using Python, Arduino Uno, USB serial, PCA9685 PWM boards, MOSFET stages, solenoid actuators, and sustain-pedal servo control."
        ),
        "Winch Lift Stand: Designed and analyzed an aluminum stand supporting a winch to raise a 150 lb motor ~50 ft with a high factor of safety.": (
            "Winch Lift Stand | August 2024: Designed and analyzed an aluminum stand supporting a winch to raise a 150 lb motor ~50 ft with a high factor of safety."
        ),
        "Food Stirrer Prototype: Performed stress analysis and supported iterative design for a 3D-printed liquid-food stirrer with a 4-person team.": (
            "Food Stirrer Prototype | Fall 2025: Performed stress analysis and supported iterative design for a 3D-printed liquid-food stirrer with a 4-person team."
        ),
        "Python Automation: Built a “Jarvis” assistant integrating calendar + smart-home control + weather/maps services; developed a versioned chess bot (~1000 Elo).": (
            "Python Automation | Summer 2024: Built a Jarvis assistant integrating calendar, smart-home, weather, and maps services; developed a versioned chess bot (~1000 Elo)."
        ),
    }

    capstone_extra = [
        "Implemented MIDI scheduling, Tkinter/CLI operator workflow, JSON hardware configuration, chunked serial transfer, 48-event ring buffer, non-blocking Arduino playback, diagnostics, calibration sweeps, and pause/resume/stop/all-off safety commands.",
    ]

    for paragraph in list(doc.paragraphs):
        if paragraph.text in replacements:
            original = paragraph.text
            set_text(paragraph, replacements[original])
            if original.startswith("Autonomous Piano Player"):
                anchor = paragraph
                for item in capstone_extra:
                    anchor = insert_paragraph_after(anchor, item)

    # Merge the coursework line back into one paragraph and remove the Google
    # Docs hard-wrapped continuation paragraph.
    for paragraph in list(doc.paragraphs):
        if paragraph.text.startswith("Relevant Coursework:"):
            set_text(
                paragraph,
                "Relevant Coursework: Space Systems Engineering, Systems Dynamics & Control, "
                "Dynamics, Heat Transfer, Mechanics of Materials, Instrumentation",
            )
        elif paragraph.text == "Transfer, Mechanics of Materials, Instrumentation":
            remove_paragraph(paragraph)

    # Remove blank paragraphs directly after these section headings.
    paragraphs = list(doc.paragraphs)
    for idx, paragraph in enumerate(paragraphs[:-1]):
        if paragraph.text.strip() in {"Education", "Professional Experience"}:
            nxt = paragraphs[idx + 1]
            if not nxt.text.strip():
                remove_paragraph(nxt)

    # Preserve the original look, but tighten pagination-sensitive spacing.
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            compact(paragraph, before=0, after=0, line=1.0)
            continue

        if text in {
            "Personal Summary",
            "Education",
            "Professional Experience",
            "Technical Skills",
            "Certifications",
            "Achievements",
            "Projects",
        }:
            compact(paragraph, before=0, after=0, line=0.9)
            for run in paragraph.runs:
                set_run_font(run, name="Times New Roman", size=12, bold=True)
        elif text.startswith(("Engineering Intern |", "Wildland Firefighter |")):
            compact(paragraph, before=0, after=0, line=0.95)
            for run in paragraph.runs:
                set_run_font(run, name="Times New Roman", size=10, bold=True)
        elif paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
            compact(paragraph, before=0, after=0, line=0.95)
            for run in paragraph.runs:
                set_run_font(run, name="Times New Roman", size=9.6)
        else:
            compact(paragraph, before=0, after=0, line=0.95)
            for run in paragraph.runs:
                set_run_font(run, name="Times New Roman", size=9.6)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
