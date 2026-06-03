from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE = Path(__file__).resolve().parent
SOURCE = BASE / "derek-martin-resume-source.docx"
OUTPUT = BASE / "derek-martin-resume-usajobs-nasa-working.docx"

BODY_FONT = "Calibri"


def set_run_font(run, size_pt=10, bold=None, italic=None):
    run.font.name = BODY_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para_spacing(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def insert_paragraph_after(paragraph, text="", style=None):
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    paragraph._p.addnext(new_para._p)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def copy_paragraph_shape(src, dst):
    if src._p.pPr is not None:
        dst._p.get_or_add_pPr().append(deepcopy(src._p.pPr))
    dst.alignment = src.alignment


def remove_numbering(paragraph):
    p_pr = paragraph._p.pPr
    if p_pr is not None and p_pr.numPr is not None:
        p_pr.remove(p_pr.numPr)


def main():
    doc = Document(SOURCE)

    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    styles = doc.styles
    styles["Normal"].font.name = BODY_FONT
    styles["Normal"].font.size = Pt(10)

    for i, paragraph in enumerate(doc.paragraphs):
        set_para_spacing(paragraph, after=0, line=1.0)
        if not paragraph.text.strip():
            continue

        for run in paragraph.runs:
            set_run_font(run, 10)

        text = paragraph.text.strip()
        if i == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if paragraph.runs:
                set_run_font(paragraph.runs[0], 14, bold=True)
                for run in paragraph.runs[1:]:
                    set_run_font(run, 10)
            set_para_spacing(paragraph, after=2)
        elif text in {
            "Personal Summary",
            "Education",
            "Professional Experience",
            "Technical Skills",
            "Certifications",
            "Achievements",
            "Projects",
        }:
            for run in paragraph.runs:
                set_run_font(run, 11.5, bold=True)
            set_para_spacing(paragraph, before=3, after=1)
        elif i in {14, 21} or text.startswith(("B.S. in ", "Oregon State University")):
            for run in paragraph.runs:
                set_run_font(run, 10, bold=run.bold, italic=run.italic)

    replacements = {
        "Mechanical Engineering student (aerospace focus, physics minor) with strength in space systems thinking and hands-on mechanical design. Built Python automation tools and currently developing Python control software (and supporting CAD/integration) for an autonomous piano-player capstone (target delivery June 2026). Designed and analyzed a high-safety-factor aluminum winch stand to lift a 150 lb motor ~50 ft. Seeking remote part-time engineering work through June 2027 and launch/space roles starting 2027.":
            "Mechanical Engineering student (aerospace focus, physics minor) with strength in space systems thinking and hands-on mechanical design. Built Python automation tools and currently developing Python control software (and supporting CAD/integration) for an autonomous piano-player capstone (target delivery June 2026). Designed and analyzed a high-safety-factor aluminum winch stand to lift a 150 lb motor ~50 ft.",
        "Engineering Intern | Evensol Inc. - EDG-L (Landfill Gas Collection) | June 2024 - September 2024":
            "Engineering Intern | Evensol Inc. - EDG-L (Landfill Gas Collection) | June 2024 - September 2024 | Hours/week: varied; average 26",
        "     Wildland Firefighter | USFS Sisters Ranger District | August 2022 - September 2023":
            "Wildland Firefighter | USFS Sisters Ranger District | August 2022 - September 2023 | Grades: GS-03 (2022), GS-04 (2023) | Hours/week: varied, 40-112",
    }
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        if original_text in replacements:
            bold = any(run.bold for run in paragraph.runs)
            paragraph.clear()
            run = paragraph.add_run(replacements[original_text])
            set_run_font(run, 10, bold=bold)
            set_para_spacing(paragraph, before=1, after=0)

    # Keep the Google Docs layout compact after adding required federal fields.
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Relevant Coursework:"):
            paragraph.text = (
                "Relevant Coursework: Space Systems Engineering, Systems Dynamics & Control, "
                "Dynamics, Heat Transfer, Mechanics of Materials, Instrumentation"
            )
            for run in paragraph.runs:
                set_run_font(run, 10)
        elif paragraph.text == "Transfer, Mechanics of Materials, Instrumentation":
            p = paragraph._element
            p.getparent().remove(p)

    date_placeholders = {
        'Awarded "Hardest Worker" in Calculus (2021)':
            'Awarded "Hardest Worker" in Calculus | Dates: 2021 and 2022 [add months if available]',
        "Team Captain for Nordic Skiing":
            "Team Captain for Nordic Skiing | Dates: 2021 - 2022 [add months if available]",
        "Medallion awarded by Bend Police Department for integrity and discipline (2021)":
            "Medallion awarded by Bend Police Department for integrity and discipline | Date: 2020 [add month if available]",
        "Autonomous Piano Player (Capstone, 2025–2026): Developing Python control software for an electromechanical system; supporting CAD/integration as needed.":
            "Autonomous Piano Player (Capstone, January 2026 - June 2026): Project leader for a bench-scale embedded mechatronics system that converts MIDI files into real-time actuator commands for a physical piano.",
        "Winch Lift Stand: Designed and analyzed an aluminum stand supporting a winch to raise a 150 lb motor ~50 ft with a high factor of safety.":
            "Winch Lift Stand | Date: August 2024: Designed and analyzed an aluminum stand supporting a winch to raise a 150 lb motor ~50 ft with a high factor of safety.",
        "Food Stirrer Prototype: Performed stress analysis and supported iterative design for a 3D-printed liquid-food stirrer with a 4-person team.":
            "Food Stirrer Prototype | Date: Fall 2025 [add month if available]: Performed stress analysis and supported iterative design for a 3D-printed liquid-food stirrer with a 4-person team.",
        "Python Automation: Built a “Jarvis” assistant integrating calendar + smart-home control + weather/maps services; developed a versioned chess bot (~1000 Elo).":
            "Python Automation | Date: Summer 2024 [add month if available]: Built a “Jarvis” assistant integrating calendar + smart-home control + weather/maps services; developed a versioned chess bot (~1000 Elo).",
    }

    for paragraph in doc.paragraphs:
        if paragraph.text in date_placeholders:
            original_project_text = paragraph.text
            paragraph.text = date_placeholders[original_project_text]
            for run in paragraph.runs:
                set_run_font(run, 10)
            if original_project_text.startswith("Autonomous Piano Player"):
                detail_1 = insert_paragraph_after(
                    paragraph,
                    "• Built Python MIDI processing and Tkinter/CLI operator workflows for tempo control, playable-range filtering, octave transposition, diagnostics, active-channel setup, and song playback.",
                )
                copy_paragraph_shape(paragraph, detail_1)
                detail_2 = insert_paragraph_after(
                    detail_1,
                    "• Implemented Arduino Uno firmware and a USB serial protocol with handshake/status polling, chunked event transfer, pause/resume/stop/all-off safety commands, non-blocking timed playback, and a 48-event ring buffer.",
                )
                copy_paragraph_shape(paragraph, detail_2)
                detail_3 = insert_paragraph_after(
                    detail_2,
                    "• Integrated 62 hardware channels across PCA9685 PWM boards, MOSFET stages, solenoid key actuators, and sustain-pedal servo; created JSON configuration, calibration sweeps, I2C diagnostics, tuning pulses, and generated calibration reports.",
                )
                copy_paragraph_shape(paragraph, detail_3)
                for detail in (detail_1, detail_2, detail_3):
                    detail.paragraph_format.left_indent = Inches(0.22)
                    detail.paragraph_format.first_line_indent = Inches(-0.12)
                    set_para_spacing(detail, after=0, line=1.0)
                    for run in detail.runs:
                        set_run_font(run, 10)

    for paragraph in doc.paragraphs:
        had_numbering = (
            paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None
        )
        if had_numbering:
            remove_numbering(paragraph)
            if paragraph.text.strip() and not paragraph.text.lstrip().startswith("•"):
                paragraph.text = f"• {paragraph.text.strip()}"
            paragraph.paragraph_format.left_indent = Inches(0.22)
            paragraph.paragraph_format.first_line_indent = Inches(-0.12)
            set_para_spacing(paragraph, after=0, line=1.0)
            for run in paragraph.runs:
                set_run_font(run, 10)

    # Google Docs inserted blank paragraphs after these headings; remove them
    # so the federal-required content stays tight and easy to scan.
    paragraphs = list(doc.paragraphs)
    for idx, paragraph in enumerate(paragraphs[:-1]):
        if paragraph.text.strip() in {"Education", "Professional Experience"}:
            following = paragraphs[idx + 1]
            if not following.text.strip():
                p = following._element
                p.getparent().remove(p)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
